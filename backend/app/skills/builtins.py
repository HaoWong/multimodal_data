"""
内置Skills实现
"""
from typing import Optional

from .base import Skill, SkillMetadata, SkillParameter, SkillResult, skill_registry
from app.services.ollama_client import ollama_client
from app.core.database import SessionLocal
from sqlalchemy import text


class ContentExtractSkill(Skill):
    """内容提取Skill - 从文件中提取文本内容、分析图片或分析视频"""

    def _define_metadata(self) -> SkillMetadata:
        return SkillMetadata(
            name="content_extract",
            description="从文档提取文本(PDF/Word/TXT)、分析图片(JPG/PNG)或分析视频(MP4/AVI等)，自动根据文件类型选择处理方式",
            parameters=[
                SkillParameter(
                    name="file_path",
                    type="string",
                    description="文件路径",
                    required=True
                ),
            ],
            return_type="string",
            examples=[
                'extract content from /path/to/document.pdf',
                'analyze image /path/to/photo.jpg',
                'analyze video /path/to/movie.mp4',
            ],
            tags=["content", "document", "extraction", "image", "video", "analysis"]
        )
    
    async def execute(self, file_path: str) -> SkillResult:
        try:
            import os
            from pathlib import Path
            from app.models.content_models import Content
            from app.core.database import SessionLocal
            
            # 调试信息
            print(f"\n[DEBUG] ContentExtract 输入路径: {file_path}")
            print(f"[DEBUG] 当前工作目录: {os.getcwd()}")
            
            # 获取backend目录路径
            current_file = Path(__file__)
            backend_dir = current_file.parent.parent.parent  # app/skills/builtins.py -> app -> backend
            uploads_dir = backend_dir / "uploads" / "contents"
            
            print(f"[DEBUG] Backend目录: {backend_dir}")
            print(f"[DEBUG] Backend目录存在: {backend_dir.exists()}")
            print(f"[DEBUG] Uploads目录: {uploads_dir}")
            print(f"[DEBUG] Uploads目录存在: {uploads_dir.exists()}")
            
            # 如果路径不存在，尝试通过数据库查找
            if not os.path.exists(file_path):
                filename = os.path.basename(file_path)
                print(f"[DEBUG] 文件不存在，尝试查找文件名: {filename}")
                
                # 首先尝试从数据库中查找
                db = SessionLocal()
                try:
                    # 通过original_name查找
                    content = db.query(Content).filter(Content.original_name == filename).first()
                    if content and content.source_path and os.path.exists(content.source_path):
                        file_path = content.source_path
                        print(f"[DEBUG] 从数据库找到文件路径: {file_path}")
                    else:
                        # 尝试查找文件名包含输入的情况
                        content = db.query(Content).filter(Content.original_name.contains(filename)).first()
                        if content and content.source_path and os.path.exists(content.source_path):
                            file_path = content.source_path
                            print(f"[DEBUG] 从数据库模糊匹配找到文件路径: {file_path}")
                finally:
                    db.close()
                
                # 如果数据库中没找到，在uploads目录中查找
                if not os.path.exists(file_path) and uploads_dir.exists():
                    print("[DEBUG] 在uploads目录中搜索...")
                    all_files = []
                    for root, dirs, files in os.walk(uploads_dir):
                        all_files.extend(files)
                        for f in files:
                            if f == filename or f.endswith(filename):
                                file_path = os.path.join(root, f)
                                print(f"[DEBUG] 找到匹配文件: {file_path}")
                                break
                        if os.path.exists(file_path):
                            break
                    print(f"[DEBUG] uploads目录中的所有文件: {all_files}")
                
                # 如果还是找不到，返回错误
                if not os.path.exists(file_path):
                    print(f"[DEBUG] 最终未找到文件: {filename}")
                    return SkillResult(
                        success=False,
                        error=f"文件不存在: {filename}"
                    )
            
            print(f"[DEBUG] 最终文件路径: {file_path}")
            print(f"[DEBUG] 文件存在: {os.path.exists(file_path)}")
            
            # 读取文件
            with open(file_path, 'rb') as f:
                content = f.read()
            
            filename = os.path.basename(file_path)
            
            # 处理ZIP文件 - 递归分析所有内容
            if filename.endswith('.zip'):
                import zipfile
                from io import BytesIO
                import tempfile
                import shutil
                
                file_list = []
                analysis_results = []
                
                # 创建临时目录解压ZIP
                temp_dir = tempfile.mkdtemp(prefix="zip_extract_")
                
                try:
                    with zipfile.ZipFile(BytesIO(content), 'r') as zf:
                        # 解压所有文件
                        zf.extractall(temp_dir)
                        
                        # 遍历解压后的文件
                        for root, dirs, files in os.walk(temp_dir):
                            for file in files:
                                file_path = os.path.join(root, file)
                                rel_path = os.path.relpath(file_path, temp_dir)
                                file_list.append(rel_path)
                                
                                # 根据文件类型进行分析
                                try:
                                    if file.endswith(('.txt', '.md', '.py', '.js', '.json', '.xml', '.html', '.css')):
                                        # 文本文件直接读取
                                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                            text = f.read()
                                            analysis_results.append(f"\n=== 文本文件: {rel_path} ===\n{text[:2000]}...")
                                    
                                    elif file.endswith('.pdf'):
                                        # PDF文件提取文本
                                        from pypdf import PdfReader
                                        reader = PdfReader(file_path)
                                        text = ""
                                        for page in reader.pages:
                                            text += page.extract_text() + "\n"
                                        analysis_results.append(f"\n=== PDF文档: {rel_path} ===\n{text[:2000]}...")
                                    
                                    elif file.endswith(('.docx', '.doc')):
                                        # Word文档提取文本
                                        from docx import Document
                                        doc = Document(file_path)
                                        text = ""
                                        for para in doc.paragraphs:
                                            text += para.text + "\n"
                                        analysis_results.append(f"\n=== Word文档: {rel_path} ===\n{text[:2000]}...")
                                    
                                    elif file.lower().endswith(('.jpg', '.jpeg', '.png', '.gif')):
                                        # 图片文件AI分析
                                        description = await ollama_client.describe_image(file_path)
                                        analysis_results.append(f"\n=== 图片: {rel_path} ===\n{description}")
                                    
                                    elif file.lower().endswith(('.mp4', '.avi', '.mov', '.wmv')):
                                        # 视频文件分析
                                        from app.utils.video_processor import video_processor
                                        import uuid
                                        content_id = str(uuid.uuid4())
                                        video_result = await video_processor.process_video(
                                            file_path, content_id, ollama_client, max_frames=5
                                        )
                                        analysis_results.append(f"\n=== 视频: {rel_path} ===\n时长: {video_result['duration']:.1f}秒\n{video_result['overall_description'][:1500]}...")
                                    
                                except Exception as e:
                                    analysis_results.append(f"\n=== 文件: {rel_path} ===\n分析失败: {str(e)}")
                
                finally:
                    # 清理临时目录
                    shutil.rmtree(temp_dir, ignore_errors=True)
                
                # 构建完整结果
                result = f"ZIP文件包含 {len(file_list)} 个文件:\n" + "\n".join(file_list)
                if analysis_results:
                    result += "\n\n=== 详细内容分析 ===" + "\n".join(analysis_results)
            
            # 处理PDF文件
            elif filename.endswith('.pdf'):
                from pypdf import PdfReader
                from io import BytesIO
                reader = PdfReader(BytesIO(content))
                text_content = ""
                for page in reader.pages:
                    text_content += page.extract_text() + "\n"
                result = text_content.strip()
            
            # 处理Word文档
            elif filename.endswith(('.docx', '.doc')):
                from docx import Document
                from io import BytesIO
                doc = Document(BytesIO(content))
                text_content = ""
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
                result = text_content.strip()
            
            # 处理图片文件 - 自动调用图像分析
            elif filename.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
                # 对于图片文件，调用图像分析
                print(f"[DEBUG] 检测到图片文件，调用图像分析: {file_path}")
                try:
                    description = await ollama_client.describe_image(file_path)
                    return SkillResult(
                        success=True,
                        data=f"图片内容分析:\n{description}",
                        metadata={"file_path": file_path, "type": "image"}
                    )
                except Exception as e:
                    return SkillResult(
                        success=False,
                        error=f"图片分析失败: {str(e)}"
                    )

            # 处理视频文件 - 自动调用视频分析
            elif filename.lower().endswith(('.mp4', '.avi', '.mov', '.wmv', '.mkv', '.flv')):
                # 对于视频文件，调用视频分析
                print(f"[DEBUG] 检测到视频文件，调用视频分析: {file_path}")
                try:
                    from app.utils.video_processor import video_processor
                    import uuid

                    content_id = str(uuid.uuid4())
                    result = await video_processor.process_video(
                        file_path,
                        content_id,
                        ollama_client,
                        max_frames=8
                    )

                    # 构建视频分析报告
                    report = f"""视频内容分析:

时长: {result['duration']:.1f}秒
分析帧数: {len(result['frames'])}

整体描述:
{result['overall_description']}

关键帧描述:"""

                    for i, frame in enumerate(result['frames'], 1):
                        report += f"\n\n帧 {i} (时间点: {frame.timestamp:.1f}s):\n{frame.description}"

                    return SkillResult(
                        success=True,
                        data=report,
                        metadata={"file_path": file_path, "type": "video", "duration": result['duration']}
                    )
                except Exception as e:
                    return SkillResult(
                        success=False,
                        error=f"视频分析失败: {str(e)}"
                    )

            else:
                # 文本文件
                result = content.decode('utf-8', errors='ignore')
            
            return SkillResult(
                success=True,
                data=result,
                metadata={"file_type": Path(file_path).suffix, "file_name": filename}
            )
        except Exception as e:
            return SkillResult(
                success=False,
                error=f"内容提取失败: {str(e)}"
            )


class ImageAnalyzeSkill(Skill):
    """图像分析Skill - 使用视觉模型分析图片"""
    
    def _define_metadata(self) -> SkillMetadata:
        return SkillMetadata(
            name="image_analyze",
            description="使用AI视觉模型分析图片内容，生成详细描述",
            parameters=[
                SkillParameter(
                    name="image_path",
                    type="string",
                    description="图片文件路径",
                    required=True
                ),
                SkillParameter(
                    name="prompt",
                    type="string",
                    description="分析提示词",
                    required=False,
                    default="请详细描述这张图片的内容"
                ),
            ],
            return_type="string",
            examples=[
                'analyze image at /path/to/image.jpg',
                'analyze image with prompt "描述图中的主要物体"',
            ],
            tags=["image", "vision", "analysis"]
        )
    
    async def execute(self, image_path: str, prompt: str = "请详细描述这张图片的内容") -> SkillResult:
        try:
            description = await ollama_client.describe_image(image_path)
            return SkillResult(
                success=True,
                data=description,
                metadata={"image_path": image_path}
            )
        except Exception as e:
            return SkillResult(
                success=False,
                error=f"图像分析失败: {str(e)}"
            )


class VideoAnalyzeSkill(Skill):
    """视频分析Skill - 提取帧并分析视频内容"""
    
    def _define_metadata(self) -> SkillMetadata:
        return SkillMetadata(
            name="video_analyze",
            description="提取视频关键帧并使用AI分析视频内容",
            parameters=[
                SkillParameter(
                    name="video_path",
                    type="string",
                    description="视频文件路径",
                    required=True
                ),
                SkillParameter(
                    name="max_frames",
                    type="integer",
                    description="最大提取帧数",
                    required=False,
                    default=8
                ),
            ],
            return_type="object",
            examples=[
                'analyze video at /path/to/video.mp4',
                'analyze video with max_frames 10',
            ],
            tags=["video", "analysis", "frames"]
        )
    
    async def execute(self, video_path: str, max_frames: int = 8) -> SkillResult:
        try:
            from app.utils.video_processor import video_processor
            import uuid
            
            content_id = str(uuid.uuid4())
            result = await video_processor.process_video(
                video_path,
                content_id,
                ollama_client,
                max_frames=max_frames
            )
            
            return SkillResult(
                success=True,
                data={
                    "description": result["overall_description"],
                    "duration": result["duration"],
                    "frame_count": len(result["frames"]),
                    "frames": [
                        {
                            "timestamp": f.timestamp,
                            "description": f.description
                        }
                        for f in result["frames"]
                    ]
                },
                metadata={"video_path": video_path}
            )
        except Exception as e:
            return SkillResult(
                success=False,
                error=f"视频分析失败: {str(e)}"
            )


class TextEmbedSkill(Skill):
    """文本向量化Skill - 生成文本的向量嵌入"""
    
    def _define_metadata(self) -> SkillMetadata:
        return SkillMetadata(
            name="text_embed",
            description="将文本转换为向量嵌入，用于语义搜索",
            parameters=[
                SkillParameter(
                    name="text",
                    type="string",
                    description="要向量化的文本",
                    required=True
                ),
            ],
            return_type="array",
            examples=[
                'embed text "这是一段示例文本"',
            ],
            tags=["embedding", "vector", "text"]
        )
    
    async def execute(self, text: str) -> SkillResult:
        try:
            embeddings = await ollama_client.embed([text])
            return SkillResult(
                success=True,
                data=embeddings[0] if embeddings else [],
                metadata={"text_length": len(text)}
            )
        except Exception as e:
            return SkillResult(
                success=False,
                error=f"文本向量化失败: {str(e)}"
            )


class VectorSearchSkill(Skill):
    """向量搜索Skill - 在RAG数据库中搜索相似内容"""
    
    def _define_metadata(self) -> SkillMetadata:
        return SkillMetadata(
            name="vector_search",
            description="使用语义搜索在RAG数据库中查找相关内容",
            parameters=[
                SkillParameter(
                    name="query",
                    type="string",
                    description="搜索查询",
                    required=True
                ),
                SkillParameter(
                    name="content_type",
                    type="string",
                    description="内容类型过滤 (text/image/video)",
                    required=False,
                    default=None
                ),
                SkillParameter(
                    name="top_k",
                    type="integer",
                    description="返回结果数量",
                    required=False,
                    default=5
                ),
            ],
            return_type="array",
            examples=[
                'search "人工智能的发展历史"',
                'search "机器学习" with content_type "text"',
            ],
            tags=["search", "rag", "vector"]
        )
    
    async def execute(self, query: str, content_type: Optional[str] = None, top_k: int = 5) -> SkillResult:
        try:
            # 生成查询向量
            embeddings = await ollama_client.embed([query])
            query_embedding = embeddings[0]
            
            # 执行向量搜索
            db = SessionLocal()
            try:
                # 构建SQL
                sql = """
                SELECT 
                    id,
                    content_type,
                    original_name,
                    extracted_text,
                    description,
                    1 - (embedding <=> :query_embedding) as similarity
                FROM contents
                WHERE 1 - (embedding <=> :query_embedding) > 0.5
                """
                
                params = {"query_embedding": str(query_embedding), "top_k": top_k}
                
                if content_type:
                    # 将字符串转换为ContentType枚举值（支持大小写）
                    from app.models.content_models import ContentType
                    try:
                        # 尝试直接匹配（大写）
                        content_type_enum = ContentType(content_type.upper())
                        sql += " AND content_type = :content_type"
                        params["content_type"] = content_type_enum.value
                    except ValueError:
                        # 如果转换失败，忽略content_type过滤
                        pass
                
                sql += " ORDER BY embedding <=> :query_embedding LIMIT :top_k"
                
                result = db.execute(text(sql), params)
                
                results = []
                for row in result:
                    results.append({
                        "id": str(row.id),
                        "content_type": row.content_type,
                        "original_name": row.original_name,
                        "text": row.extracted_text or row.description,
                        "similarity": row.similarity
                    })
                
                return SkillResult(
                    success=True,
                    data=results,
                    metadata={"query": query, "result_count": len(results)}
                )
            finally:
                db.close()
                
        except Exception as e:
            return SkillResult(
                success=False,
                error=f"向量搜索失败: {str(e)}"
            )


class LLMChatSkill(Skill):
    """LLM对话Skill - 与语言模型对话"""
    
    def _define_metadata(self) -> SkillMetadata:
        return SkillMetadata(
            name="llm_chat",
            description="与大型语言模型进行对话",
            parameters=[
                SkillParameter(
                    name="message",
                    type="string",
                    description="用户消息",
                    required=True
                ),
                SkillParameter(
                    name="system_prompt",
                    type="string",
                    description="系统提示词",
                    required=False,
                    default="你是一个有帮助的助手"
                ),
                SkillParameter(
                    name="context",
                    type="array",
                    description="对话上下文",
                    required=False,
                    default=[]
                ),
            ],
            return_type="string",
            examples=[
                'chat "你好，请介绍一下自己"',
                'chat with system_prompt "你是专家"',
            ],
            tags=["llm", "chat", "generation"]
        )
    
    async def execute(self, message: str, system_prompt: str = "你是一个有帮助的助手", context: list = None) -> SkillResult:
        try:
            # 构建消息
            messages = []
            if system_prompt:
                messages.append({"role": "system", "content": system_prompt})
            
            # 添加上下文
            if context:
                messages.extend(context)
            
            messages.append({"role": "user", "content": message})
            
            # 调用LLM
            response_text = ""
            async for chunk in ollama_client.chat(messages, stream=False):
                response_text += chunk
            
            return SkillResult(
                success=True,
                data=response_text,
                metadata={"model": ollama_client.chat_model}
            )
        except Exception as e:
            return SkillResult(
                success=False,
                error=f"LLM对话失败: {str(e)}"
            )


# 注册所有内置Skills
def register_builtin_skills():
    """注册所有内置Skills"""
    skill_registry.register(ContentExtractSkill())
    skill_registry.register(ImageAnalyzeSkill())
    skill_registry.register(VideoAnalyzeSkill())
    skill_registry.register(TextEmbedSkill())
    skill_registry.register(VectorSearchSkill())
    skill_registry.register(LLMChatSkill())
    print("✅ All builtin skills registered")
