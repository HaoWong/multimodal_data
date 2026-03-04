"""
统一文件处理器模块
提供抽象基类和具体实现，支持多种文件类型的处理
"""
import os
import io
import shutil
import zipfile
import tempfile
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, BinaryIO
from dataclasses import dataclass, field
from enum import Enum
import mimetypes

from fastapi import UploadFile


class FileType(Enum):
    """文件类型枚举"""
    TEXT = "text"
    PDF = "pdf"
    WORD = "word"
    IMAGE = "image"
    VIDEO = "video"
    ZIP = "zip"
    UNKNOWN = "unknown"


class FileProcessError(Exception):
    """文件处理错误"""
    pass


class FileValidationError(FileProcessError):
    """文件验证错误"""
    pass


class FileExtractError(FileProcessError):
    """内容提取错误"""
    pass


@dataclass
class FileMetadata:
    """文件元数据"""
    filename: str
    file_type: FileType
    mime_type: str
    size: int
    extension: str
    extra: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ProcessResult:
    """处理结果"""
    success: bool
    content: Optional[str] = None
    embedding: Optional[List[float]] = None
    metadata: Optional[FileMetadata] = None
    chunks: List[Dict] = field(default_factory=list)
    error_message: Optional[str] = None
    extra_data: Dict[str, Any] = field(default_factory=dict)


class FileStorageManager:
    """统一文件存储路径管理器"""
    
    def __init__(self, base_dir: str = "uploads"):
        self.base_dir = Path(base_dir)
        self._ensure_directories()
    
    def _ensure_directories(self):
        """确保所有必要的目录存在"""
        directories = [
            self.base_dir / "files",
            self.base_dir / "frames",
            self.base_dir / "temp",
            self.base_dir / "thumbnails",
            self.base_dir / "extracted",
        ]
        for directory in directories:
            directory.mkdir(parents=True, exist_ok=True)
    
    def get_file_path(self, content_id: str, filename: str) -> Path:
        """获取文件存储路径"""
        return self.base_dir / "files" / f"{content_id}_{filename}"
    
    def get_frames_dir(self, content_id: str) -> Path:
        """获取视频帧存储目录"""
        frames_dir = self.base_dir / "frames" / content_id
        frames_dir.mkdir(parents=True, exist_ok=True)
        return frames_dir
    
    def get_temp_dir(self, content_id: str) -> Path:
        """获取临时目录"""
        temp_dir = self.base_dir / "temp" / content_id
        temp_dir.mkdir(parents=True, exist_ok=True)
        return temp_dir
    
    def get_thumbnail_path(self, content_id: str, ext: str = "jpg") -> Path:
        """获取缩略图路径"""
        return self.base_dir / "thumbnails" / f"{content_id}.{ext}"
    
    def get_extracted_dir(self, content_id: str) -> Path:
        """获取解压文件目录"""
        extracted_dir = self.base_dir / "extracted" / content_id
        extracted_dir.mkdir(parents=True, exist_ok=True)
        return extracted_dir
    
    def cleanup(self, content_id: str):
        """清理与content_id相关的所有临时文件"""
        paths_to_clean = [
            self.base_dir / "frames" / content_id,
            self.base_dir / "temp" / content_id,
            self.base_dir / "extracted" / content_id,
        ]
        for path in paths_to_clean:
            if path.exists():
                if path.is_dir():
                    shutil.rmtree(path)
                else:
                    path.unlink()


class FileProcessor(ABC):
    """文件处理器抽象基类"""
    
    # 支持的文件扩展名
    SUPPORTED_EXTENSIONS: Tuple[str, ...] = ()
    
    # 文件类型
    FILE_TYPE: FileType = FileType.UNKNOWN
    
    def __init__(
        self,
        storage_manager: Optional[FileStorageManager] = None,
        ollama_client=None
    ):
        self.storage = storage_manager or FileStorageManager()
        self.ollama_client = ollama_client
    
    @abstractmethod
    async def validate(self, file_content: bytes, filename: str) -> bool:
        """
        验证文件是否有效
        
        Args:
            file_content: 文件内容字节
            filename: 文件名
            
        Returns:
            验证是否通过
            
        Raises:
            FileValidationError: 验证失败时抛出
        """
        pass
    
    @abstractmethod
    async def extract_content(
        self,
        file_content: bytes,
        filename: str,
        content_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        提取文件内容
        
        Args:
            file_content: 文件内容字节
            filename: 文件名
            content_id: 内容ID（用于存储临时文件）
            
        Returns:
            包含提取内容的字典
            
        Raises:
            FileExtractError: 提取失败时抛出
        """
        pass
    
    async def generate_embedding(self, text: str) -> List[float]:
        """
        生成文本的嵌入向量
        
        Args:
            text: 要生成嵌入的文本
            
        Returns:
            嵌入向量
        """
        if not self.ollama_client:
            raise FileProcessError("Ollama client not configured")
        
        from app.utils.text_chunker import BatchEmbedder
        
        result = await BatchEmbedder.embed_large_text(
            text,
            self.ollama_client._embed_single_batch
        )
        return result.get("embedding", [])
    
    async def process(
        self,
        file_content: bytes,
        filename: str,
        content_id: Optional[str] = None,
        generate_embedding: bool = True
    ) -> ProcessResult:
        """
        处理文件的完整流程
        
        Args:
            file_content: 文件内容字节
            filename: 文件名
            content_id: 内容ID
            generate_embedding: 是否生成嵌入向量
            
        Returns:
            ProcessResult: 处理结果
        """
        try:
            # 1. 验证文件
            await self.validate(file_content, filename)
            
            # 2. 提取内容
            extract_result = await self.extract_content(
                file_content, filename, content_id
            )
            
            content = extract_result.get("content", "")
            metadata_dict = extract_result.get("metadata", {})
            extra_data = extract_result.get("extra_data", {})
            chunks = extract_result.get("chunks", [])
            
            # 3. 创建元数据
            metadata = FileMetadata(
                filename=filename,
                file_type=self.FILE_TYPE,
                mime_type=mimetypes.guess_type(filename)[0] or "application/octet-stream",
                size=len(file_content),
                extension=Path(filename).suffix.lower(),
                extra=metadata_dict
            )
            
            # 4. 生成嵌入向量
            embedding = None
            if generate_embedding and content:
                try:
                    embedding = await self.generate_embedding(content)
                except Exception as e:
                    extra_data["embedding_error"] = str(e)
            
            return ProcessResult(
                success=True,
                content=content,
                embedding=embedding,
                metadata=metadata,
                chunks=chunks,
                extra_data=extra_data
            )
            
        except FileValidationError as e:
            return ProcessResult(
                success=False,
                error_message=f"验证失败: {str(e)}"
            )
        except FileExtractError as e:
            return ProcessResult(
                success=False,
                error_message=f"内容提取失败: {str(e)}"
            )
        except Exception as e:
            return ProcessResult(
                success=False,
                error_message=f"处理失败: {str(e)}"
            )
    
    @classmethod
    def get_file_type(cls) -> FileType:
        """获取文件类型"""
        return cls.FILE_TYPE
    
    @classmethod
    def supports_extension(cls, extension: str) -> bool:
        """检查是否支持指定扩展名"""
        ext = extension.lower()
        if not ext.startswith('.'):
            ext = '.' + ext
        return ext in cls.SUPPORTED_EXTENSIONS


class TextFileProcessor(FileProcessor):
    """文本文件处理器"""
    
    SUPPORTED_EXTENSIONS = ('.txt', '.md', '.json', '.xml', '.csv', '.log', '.py', '.js', '.html', '.css')
    FILE_TYPE = FileType.TEXT
    
    # 最大文件大小 (10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    async def validate(self, file_content: bytes, filename: str) -> bool:
        """验证文本文件"""
        if len(file_content) > self.MAX_FILE_SIZE:
            raise FileValidationError(f"文件大小超过限制: {self.MAX_FILE_SIZE / 1024 / 1024}MB")
        
        # 尝试解码为文本
        try:
            file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                file_content.decode('gbk')
            except UnicodeDecodeError:
                raise FileValidationError("无法将文件解码为文本格式")
        
        return True
    
    async def extract_content(
        self,
        file_content: bytes,
        filename: str,
        content_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """提取文本内容"""
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            text = file_content.decode('gbk', errors='ignore')
        
        return {
            "content": text.strip(),
            "metadata": {"line_count": text.count('\n') + 1},
            "extra_data": {},
            "chunks": []
        }


class PDFFileProcessor(FileProcessor):
    """PDF文件处理器"""
    
    SUPPORTED_EXTENSIONS = ('.pdf',)
    FILE_TYPE = FileType.PDF
    
    # 最大文件大小 (50MB)
    MAX_FILE_SIZE = 50 * 1024 * 1024
    
    async def validate(self, file_content: bytes, filename: str) -> bool:
        """验证PDF文件"""
        if len(file_content) > self.MAX_FILE_SIZE:
            raise FileValidationError(f"文件大小超过限制: {self.MAX_FILE_SIZE / 1024 / 1024}MB")
        
        # 检查PDF文件头
        if not file_content.startswith(b'%PDF'):
            raise FileValidationError("无效的PDF文件格式")
        
        return True
    
    async def extract_content(
        self,
        file_content: bytes,
        filename: str,
        content_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """提取PDF内容"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(io.BytesIO(file_content))
            text_parts = []
            
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    text_parts.append(f"[第{i+1}页提取失败: {str(e)}]")
            
            full_text = "\n".join(text_parts)
            
            return {
                "content": full_text.strip(),
                "metadata": {
                    "page_count": len(reader.pages),
                    "pdf_version": reader.pdf_header if hasattr(reader, 'pdf_header') else 'unknown'
                },
                "extra_data": {},
                "chunks": []
            }
            
        except ImportError:
            raise FileExtractError("缺少pypdf库，请安装: pip install pypdf")
        except Exception as e:
            raise FileExtractError(f"PDF解析失败: {str(e)}")


class WordFileProcessor(FileProcessor):
    """Word文档处理器"""
    
    SUPPORTED_EXTENSIONS = ('.docx', '.doc')
    FILE_TYPE = FileType.WORD
    
    # 最大文件大小 (50MB)
    MAX_FILE_SIZE = 50 * 1024 * 1024
    
    async def validate(self, file_content: bytes, filename: str) -> bool:
        """验证Word文件"""
        if len(file_content) > self.MAX_FILE_SIZE:
            raise FileValidationError(f"文件大小超过限制: {self.MAX_FILE_SIZE / 1024 / 1024}MB")
        
        # 检查DOCX文件头 (ZIP格式)
        if filename.lower().endswith('.docx'):
            if not file_content.startswith(b'PK'):
                raise FileValidationError("无效的DOCX文件格式")
        
        return True
    
    async def extract_content(
        self,
        file_content: bytes,
        filename: str,
        content_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """提取Word文档内容"""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_content))
            
            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格文本
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    tables_text.append(' | '.join(row_text))
            
            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n[表格内容]\n" + "\n".join(tables_text)
            
            return {
                "content": full_text.strip(),
                "metadata": {
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables)
                },
                "extra_data": {},
                "chunks": []
            }
            
        except ImportError:
            raise FileExtractError("缺少python-docx库，请安装: pip install python-docx")
        except Exception as e:
            raise FileExtractError(f"Word文档解析失败: {str(e)}")


class ImageFileProcessor(FileProcessor):
    """图片文件处理器"""
    
    SUPPORTED_EXTENSIONS = ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')
    FILE_TYPE = FileType.IMAGE
    
    # 最大文件大小 (20MB)
    MAX_FILE_SIZE = 20 * 1024 * 1024
    
    async def validate(self, file_content: bytes, filename: str) -> bool:
        """验证图片文件"""
        if len(file_content) > self.MAX_FILE_SIZE:
            raise FileValidationError(f"文件大小超过限制: {self.MAX_FILE_SIZE / 1024 / 1024}MB")
        
        # 检查常见图片格式头
        image_headers = {
            b'\xff\xd8\xff': ['.jpg', '.jpeg'],
            b'\x89PNG': ['.png'],
            b'GIF87a': ['.gif'],
            b'GIF89a': ['.gif'],
            b'BM': ['.bmp'],
            b'RIFF': ['.webp'],  # WebP starts with RIFF
        }
        
        valid = False
        for header, exts in image_headers.items():
            if file_content.startswith(header):
                valid = True
                break
        
        if not valid:
            raise FileValidationError("无效的图片文件格式")
        
        return True
    
    async def extract_content(
        self,
        file_content: bytes,
        filename: str,
        content_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """提取图片内容（生成描述）"""
        if not self.ollama_client:
            raise FileExtractError("需要Ollama客户端来生成图片描述")
        
        try:
            from PIL import Image
            
            # 保存临时文件
            temp_path = None
            if content_id:
                temp_path = self.storage.get_file_path(content_id, filename)
            else:
                temp_fd, temp_path = tempfile.mkstemp(suffix=Path(filename).suffix)
                temp_path = Path(temp_path)
                os.close(temp_fd)
            
            temp_path = Path(temp_path)
            temp_path.write_bytes(file_content)
            
            # 获取图片信息
            with Image.open(temp_path) as img:
                width, height = img.size
                format_type = img.format
                mode = img.mode
            
            # 生成图片描述
            description = await self.ollama_client.describe_image(str(temp_path))
            
            # 清理临时文件
            if not content_id and temp_path.exists():
                temp_path.unlink()
            
            return {
                "content": description,
                "metadata": {
                    "width": width,
                    "height": height,
                    "format": format_type,
                    "mode": mode
                },
                "extra_data": {
                    "image_path": str(temp_path) if content_id else None
                },
                "chunks": []
            }
            
        except ImportError:
            raise FileExtractError("缺少PIL库，请安装: pip install Pillow")
        except Exception as e:
            raise FileExtractError(f"图片处理失败: {str(e)}")


class VideoFileProcessor(FileProcessor):
    """视频文件处理器"""
    
    SUPPORTED_EXTENSIONS = ('.mp4', '.avi', '.mov', '.mkv', '.flv', '.wmv', '.webm')
    FILE_TYPE = FileType.VIDEO
    
    # 最大文件大小 (500MB)
    MAX_FILE_SIZE = 500 * 1024 * 1024
    
    # 默认提取帧数
    DEFAULT_MAX_FRAMES = 10
    
    async def validate(self, file_content: bytes, filename: str) -> bool:
        """验证视频文件"""
        if len(file_content) > self.MAX_FILE_SIZE:
            raise FileValidationError(f"文件大小超过限制: {self.MAX_FILE_SIZE / 1024 / 1024}MB")
        
        return True
    
    async def extract_content(
        self,
        file_content: bytes,
        filename: str,
        content_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """提取视频内容（帧提取+描述）"""
        if not content_id:
            raise FileExtractError("视频处理需要提供content_id")
        
        if not self.ollama_client:
            raise FileExtractError("需要Ollama客户端来生成视频描述")
        
        try:
            import cv2
            
            # 保存视频文件
            video_path = self.storage.get_file_path(content_id, filename)
            video_path.write_bytes(file_content)
            
            # 打开视频
            cap = cv2.VideoCapture(str(video_path))
            if not cap.isOpened():
                raise FileExtractError(f"无法打开视频: {filename}")
            
            # 获取视频信息
            fps = cap.get(cv2.CAP_PROP_FPS)
            total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            duration = total_frames / fps if fps > 0 else 0
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            
            # 提取帧
            frames_dir = self.storage.get_frames_dir(content_id)
            frame_descriptions = []
            frame_infos = []
            
            max_frames = self.DEFAULT_MAX_FRAMES
            frame_interval = max(2.0, duration / max_frames)  # 最小间隔2秒
            
            current_time = 0
            frame_count = 0
            
            while current_time < duration and frame_count < max_frames:
                cap.set(cv2.CAP_PROP_POS_MSEC, current_time * 1000)
                ret, frame = cap.read()
                
                if not ret:
                    break
                
                # 保存帧
                frame_filename = f"frame_{frame_count:04d}.jpg"
                frame_path = frames_dir / frame_filename
                cv2.imwrite(str(frame_path), frame)
                
                # 生成帧描述
                try:
                    description = await self.ollama_client.describe_image(str(frame_path))
                    frame_descriptions.append(f"[{current_time:.1f}s] {description}")
                    frame_infos.append({
                        "timestamp": current_time,
                        "frame_number": frame_count,
                        "path": str(frame_path),
                        "description": description
                    })
                except Exception as e:
                    frame_descriptions.append(f"[{current_time:.1f}s] 描述生成失败")
                    frame_infos.append({
                        "timestamp": current_time,
                        "frame_number": frame_count,
                        "path": str(frame_path),
                        "description": f"错误: {str(e)}"
                    })
                
                current_time += frame_interval
                frame_count += 1
            
            cap.release()
            
            # 生成整体描述
            overall_description = "\n".join(frame_descriptions)
            
            return {
                "content": overall_description,
                "metadata": {
                    "duration": duration,
                    "fps": fps,
                    "width": width,
                    "height": height,
                    "total_frames": total_frames,
                    "extracted_frames": frame_count
                },
                "extra_data": {
                    "frames_dir": str(frames_dir),
                    "video_path": str(video_path),
                    "frame_infos": frame_infos
                },
                "chunks": []
            }
            
        except ImportError:
            raise FileExtractError("缺少cv2库，请安装: pip install opencv-python")
        except Exception as e:
            raise FileExtractError(f"视频处理失败: {str(e)}")


class ZIPFileProcessor(FileProcessor):
    """ZIP压缩包处理器"""
    
    SUPPORTED_EXTENSIONS = ('.zip',)
    FILE_TYPE = FileType.ZIP
    
    # 最大文件大小 (100MB)
    MAX_FILE_SIZE = 100 * 1024 * 1024
    
    async def validate(self, file_content: bytes, filename: str) -> bool:
        """验证ZIP文件"""
        if len(file_content) > self.MAX_FILE_SIZE:
            raise FileValidationError(f"文件大小超过限制: {self.MAX_FILE_SIZE / 1024 / 1024}MB")
        
        # 检查ZIP文件头
        if not file_content.startswith(b'PK'):
            raise FileValidationError("无效的ZIP文件格式")
        
        # 尝试打开验证
        try:
            with zipfile.ZipFile(io.BytesIO(file_content), 'r') as zf:
                # 检查是否有损坏的文件
                bad_file = zf.testzip()
                if bad_file:
                    raise FileValidationError(f"ZIP文件包含损坏的文件: {bad_file}")
        except zipfile.BadZipFile:
            raise FileValidationError("无效的ZIP文件")
        
        return True
    
    async def extract_content(
        self,
        file_content: bytes,
        filename: str,
        content_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """提取ZIP内容（列出文件并尝试提取文本文件内容）"""
        try:
            extracted_contents = []
            file_list = []
            
            with zipfile.ZipFile(io.BytesIO(file_content), 'r') as zf:
                for info in zf.infolist():
                    file_list.append({
                        "name": info.filename,
                        "size": info.file_size,
                        "compressed_size": info.compress_size,
                        "is_dir": info.is_dir()
                    })
                    
                    # 尝试提取文本文件内容
                    if not info.is_dir() and info.file_size < 1024 * 1024:  # 只处理小于1MB的文件
                        try:
                            with zf.open(info.filename) as f:
                                file_bytes = f.read()
                                # 尝试解码为文本
                                try:
                                    text = file_bytes.decode('utf-8')
                                    extracted_contents.append(f"=== {info.filename} ===\n{text[:2000]}")
                                except:
                                    pass  # 忽略无法解码的文件
                        except:
                            pass
            
            content = f"ZIP文件包含 {len(file_list)} 个条目\n\n"
            content += "文件列表:\n"
            for f in file_list:
                content += f"  {'[DIR] ' if f['is_dir'] else '[FILE]' } {f['name']} ({f['size']} bytes)\n"
            
            if extracted_contents:
                content += "\n\n提取的文本内容:\n" + "\n\n".join(extracted_contents)
            
            return {
                "content": content,
                "metadata": {
                    "entry_count": len(file_list),
                    "file_count": sum(1 for f in file_list if not f['is_dir']),
                    "dir_count": sum(1 for f in file_list if f['is_dir'])
                },
                "extra_data": {
                    "file_list": file_list
                },
                "chunks": []
            }
            
        except Exception as e:
            raise FileExtractError(f"ZIP处理失败: {str(e)}")


class FileProcessorFactory:
    """文件处理器工厂类"""
    
    # 注册所有处理器
    _processors: List[type] = [
        TextFileProcessor,
        PDFFileProcessor,
        WordFileProcessor,
        ImageFileProcessor,
        VideoFileProcessor,
        ZIPFileProcessor,
    ]
    
    def __init__(
        self,
        storage_manager: Optional[FileStorageManager] = None,
        ollama_client=None
    ):
        self.storage = storage_manager or FileStorageManager()
        self.ollama_client = ollama_client
        self._processor_instances: Dict[FileType, FileProcessor] = {}
    
    def get_processor(self, filename: str) -> Optional[FileProcessor]:
        """
        根据文件名获取对应的处理器
        
        Args:
            filename: 文件名
            
        Returns:
            对应的文件处理器实例，如果没有找到则返回None
        """
        ext = Path(filename).suffix.lower()
        
        for processor_class in self._processors:
            if processor_class.supports_extension(ext):
                file_type = processor_class.get_file_type()
                
                # 缓存处理器实例
                if file_type not in self._processor_instances:
                    self._processor_instances[file_type] = processor_class(
                        storage_manager=self.storage,
                        ollama_client=self.ollama_client
                    )
                
                return self._processor_instances[file_type]
        
        return None
    
    def get_processor_by_type(self, file_type: FileType) -> Optional[FileProcessor]:
        """
        根据文件类型获取处理器
        
        Args:
            file_type: 文件类型
            
        Returns:
            对应的文件处理器实例
        """
        if file_type in self._processor_instances:
            return self._processor_instances[file_type]
        
        for processor_class in self._processors:
            if processor_class.get_file_type() == file_type:
                instance = processor_class(
                    storage_manager=self.storage,
                    ollama_client=self.ollama_client
                )
                self._processor_instances[file_type] = instance
                return instance
        
        return None
    
    async def process_file(
        self,
        file_content: bytes,
        filename: str,
        content_id: Optional[str] = None,
        generate_embedding: bool = True
    ) -> ProcessResult:
        """
        自动选择处理器并处理文件
        
        Args:
            file_content: 文件内容字节
            filename: 文件名
            content_id: 内容ID
            generate_embedding: 是否生成嵌入向量
            
        Returns:
            ProcessResult: 处理结果
        """
        processor = self.get_processor(filename)
        
        if not processor:
            return ProcessResult(
                success=False,
                error_message=f"不支持的文件格式: {Path(filename).suffix}"
            )
        
        return await processor.process(
            file_content,
            filename,
            content_id,
            generate_embedding
        )
    
    @classmethod
    def register_processor(cls, processor_class: type):
        """注册新的处理器类"""
        if not issubclass(processor_class, FileProcessor):
            raise ValueError("处理器必须继承自FileProcessor")
        
        if processor_class not in cls._processors:
            cls._processors.append(processor_class)
    
    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        """获取所有支持的文件扩展名"""
        extensions = []
        for processor_class in cls._processors:
            extensions.extend(processor_class.SUPPORTED_EXTENSIONS)
        return extensions
    
    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """检查是否支持指定文件"""
        ext = Path(filename).suffix.lower()
        return any(
            processor_class.supports_extension(ext)
            for processor_class in cls._processors
        )


# 便捷函数
async def process_upload_file(
    upload_file: UploadFile,
    content_id: Optional[str] = None,
    ollama_client=None,
    generate_embedding: bool = True
) -> ProcessResult:
    """
    处理上传的文件
    
    Args:
        upload_file: FastAPI上传文件对象
        content_id: 内容ID
        ollama_client: Ollama客户端
        generate_embedding: 是否生成嵌入向量
        
    Returns:
        ProcessResult: 处理结果
    """
    content = await upload_file.read()
    filename = upload_file.filename
    
    factory = FileProcessorFactory(ollama_client=ollama_client)
    return await factory.process_file(
        content,
        filename,
        content_id,
        generate_embedding
    )


# 全局工厂实例
def get_file_processor_factory(ollama_client=None) -> FileProcessorFactory:
    """获取文件处理器工厂实例"""
    return FileProcessorFactory(ollama_client=ollama_client)
