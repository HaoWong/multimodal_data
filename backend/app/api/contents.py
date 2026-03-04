"""
统一内容管理API
整合文档、图片、视频等多种内容类型的管理接口
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query
from fastapi.responses import RedirectResponse
from sqlalchemy.orm import Session
from typing import Optional, List
from uuid import UUID
import os
import shutil
import uuid

from app.core.database import get_db
from app.core.task_manager import task_manager
from app.core.response import APIResponse, ErrorCode
from app.models.content_models import Content, ContentType
from app.models.schemas import BaseResponse, DocumentCreate, DocumentResponse, SearchRequest, SearchResponse
from app.models.database_models import Image
from app.services.ollama_client import ollama_client
from app.services.document_service import DocumentService
from app.utils.file_parser import parse_file

router = APIRouter(prefix="/contents", tags=["内容管理"])

UPLOAD_DIR = "uploads/contents"
IMAGE_UPLOAD_DIR = "uploads/images"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(IMAGE_UPLOAD_DIR, exist_ok=True)


def get_content_type(filename: str) -> ContentType:
    """根据文件名判断内容类型"""
    ext = os.path.splitext(filename)[1].lower()
    if ext in ['.txt', '.pdf', '.docx', '.doc', '.md']:
        return ContentType.TEXT
    elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
        return ContentType.IMAGE
    elif ext in ['.mp4', '.avi', '.mov', '.wmv', '.mkv']:
        return ContentType.VIDEO
    elif ext in ['.zip']:
        return ContentType.TEXT  # ZIP作为文本类型处理，但会特殊处理
    else:
        return ContentType.TEXT


# ==================== 统一内容管理端点 ====================

@router.post("/upload")
async def upload_content(
    file: UploadFile = File(...),
    metadata: Optional[str] = Form(None),
    chunk_size: Optional[int] = Form(1000, description="文本分块大小（字符数），默认1000"),
    chunk_overlap: Optional[int] = Form(200, description="分块重叠大小（字符数），默认200"),
    enable_chunking: Optional[bool] = Form(True, description="是否启用文本分块，默认启用"),
    db: Session = Depends(get_db)
):
    """上传内容（文档/图片/视频）
    
    参数:
    - file: 上传的文件
    - metadata: 元数据JSON字符串
    - chunk_size: 文本分块大小（默认1000字符）
    - chunk_overlap: 分块重叠大小（默认200字符）
    - enable_chunking: 是否启用文本分块
    """
    # 生成UUID（保持为UUID对象，不要转字符串）
    content_id = uuid.uuid4()
    file_ext = os.path.splitext(file.filename)[1]
    save_path = os.path.join(UPLOAD_DIR, f"{content_id}{file_ext}")

    # 保存文件
    with open(save_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # 确定内容类型
    content_type = get_content_type(file.filename)

    # 创建内容记录
    content = Content(
        id=content_id,
        content_type=content_type,
        source_path=save_path,
        original_name=file.filename,
        file_size=os.path.getsize(save_path),
        mime_type=file.content_type,
        content_metadata={"upload_meta": metadata} if metadata else {}
    )
    
    # 提取内容
    if content_type == ContentType.TEXT:
        # 提取文本
        try:
            # 处理ZIP文件 - 递归分析所有内容
            if file.filename.endswith('.zip'):
                import zipfile
                from io import BytesIO
                import tempfile
                
                with open(save_path, 'rb') as f:
                    zip_content = f.read()
                
                # 创建临时目录解压
                temp_dir = tempfile.mkdtemp(prefix="zip_content_")
                file_list = []
                analysis_results = []
                
                try:
                    with zipfile.ZipFile(BytesIO(zip_content), 'r') as zf:
                        # 解压所有文件，保持原始文件名编码
                        for member in zf.namelist():
                            # 处理文件名编码问题
                            try:
                                # 尝试使用 UTF-8
                                member_name = member.encode('cp437').decode('utf-8')
                            except (UnicodeEncodeError, UnicodeDecodeError):
                                try:
                                    # 尝试使用 GBK（中文 Windows 常见）
                                    member_name = member.encode('cp437').decode('gbk')
                                except (UnicodeEncodeError, UnicodeDecodeError):
                                    # 使用原始名称
                                    member_name = member
                            
                            # 解压文件
                            zf.extract(member, temp_dir)
                            # 如果文件名被解码，重命名为正确的中文名
                            original_path = os.path.join(temp_dir, member)
                            decoded_path = os.path.join(temp_dir, member_name)
                            if original_path != decoded_path and os.path.exists(original_path):
                                os.rename(original_path, decoded_path)
                        
                        # 遍历解压后的文件
                        for root, dirs, files in os.walk(temp_dir):
                            for extracted_file in files:
                                file_path = os.path.join(root, extracted_file)
                                rel_path = os.path.relpath(file_path, temp_dir)
                                # 确保路径使用正确的分隔符和编码
                                rel_path = rel_path.replace('\\', '/')
                                file_list.append(rel_path)
                                
                                # 根据文件类型进行分析
                                try:
                                    if extracted_file.endswith(('.txt', '.md', '.py', '.js', '.json', '.xml', '.html', '.css')):
                                        # 文本文件直接读取
                                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                            text = f.read()
                                            analysis_results.append(f"\n=== 文本文件: {rel_path} ===\n{text[:2000]}...")
                                    
                                    elif extracted_file.endswith('.pdf'):
                                        # PDF文件提取文本
                                        from pypdf import PdfReader
                                        reader = PdfReader(file_path)
                                        text = ""
                                        for page in reader.pages:
                                            text += page.extract_text() + "\n"
                                        analysis_results.append(f"\n=== PDF文档: {rel_path} ===\n{text[:2000]}...")
                                    
                                    elif extracted_file.endswith(('.docx', '.doc')):
                                        # Word文档提取文本
                                        from docx import Document
                                        doc = Document(file_path)
                                        text = ""
                                        for para in doc.paragraphs:
                                            text += para.text + "\n"
                                        analysis_results.append(f"\n=== Word文档: {rel_path} ===\n{text[:2000]}...")
                                    
                                    elif extracted_file.lower().endswith(('.jpg', '.jpeg', '.png', '.gif')):
                                        # 图片文件AI分析
                                        description = await ollama_client.describe_image(file_path)
                                        analysis_results.append(f"\n=== 图片: {rel_path} ===\n{description}")
                                    
                                    elif extracted_file.lower().endswith(('.mp4', '.avi', '.mov', '.wmv')):
                                        # 视频文件分析
                                        from app.utils.video_processor import video_processor
                                        import uuid as uuid_module
                                        video_content_id = str(uuid_module.uuid4())
                                        result = await video_processor.process_video(
                                            file_path, video_content_id, ollama_client, max_frames=5
                                        )
                                        analysis_results.append(f"\n=== 视频: {rel_path} ===\n时长: {result['duration']:.1f}秒\n{result['overall_description'][:1500]}...")
                                    
                                except Exception as e:
                                    analysis_results.append(f"\n=== 文件: {rel_path} ===\n分析失败: {str(e)}")
                
                finally:
                    # 清理临时目录
                    shutil.rmtree(temp_dir, ignore_errors=True)
                
                # 构建完整结果
                content.extracted_text = f"ZIP文件包含 {len(file_list)} 个文件:\n" + "\n".join(file_list)
                if analysis_results:
                    content.extracted_text += "\n\n=== 详细内容分析 ===" + "\n".join(analysis_results)
                
                content.content_metadata["file_count"] = len(file_list)
                content.content_metadata["file_list"] = file_list
            
            # 处理PDF文件
            elif file.filename.endswith('.pdf'):
                from pypdf import PdfReader
                from io import BytesIO
                
                with open(save_path, 'rb') as f:
                    file_content = f.read()
                
                reader = PdfReader(BytesIO(file_content))
                text_content = ""
                for page in reader.pages:
                    text_content += page.extract_text() + "\n"
                content.extracted_text = text_content.strip()
            
            # 处理Word文档
            elif file.filename.endswith(('.docx', '.doc')):
                from docx import Document
                from io import BytesIO
                
                with open(save_path, 'rb') as f:
                    file_content = f.read()
                
                doc = Document(BytesIO(file_content))
                text_content = ""
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
                content.extracted_text = text_content.strip()
            
            # 普通文本文件
            else:
                with open(save_path, 'rb') as f:
                    file_content = f.read()
                content.extracted_text = file_content.decode('utf-8', errors='ignore')
            
            # 文本分块处理
            if enable_chunking and content.extracted_text:
                from app.utils.text_chunker import TextChunker
                
                # 使用类方法进行文本分块
                chunks = TextChunker.split_by_paragraphs(
                    content.extracted_text,
                    max_chunk_size=chunk_size
                )
                
                # 保存分块信息到元数据
                content.content_metadata["chunking"] = {
                    "enabled": True,
                    "chunk_size": chunk_size,
                    "chunk_overlap": chunk_overlap,
                    "total_chunks": len(chunks)
                }
                
                # 使用第一个分块生成向量（或者可以生成所有分块的向量）
                if chunks:
                    embeddings = await ollama_client.embed([chunks[0][:2000]])
                    content.embedding = embeddings[0] if embeddings else None
            else:
                # 不分块，直接使用全文
                if content.extracted_text:
                    embeddings = await ollama_client.embed([content.extracted_text[:2000]])
                    content.embedding = embeddings[0] if embeddings else None
                    
                content.content_metadata["chunking"] = {
                    "enabled": False
                }
        except Exception as e:
            content.extracted_text = f"提取失败: {str(e)}"
    
    elif content_type == ContentType.IMAGE:
        # 分析图片
        try:
            description = await ollama_client.describe_image(save_path)
            content.description = description
            
            # 生成向量
            embeddings = await ollama_client.embed([description])
            content.embedding = embeddings[0] if embeddings else None
        except Exception as e:
            content.description = f"分析失败: {str(e)}"
    
    elif content_type == ContentType.VIDEO:
        # 分析视频
        try:
            from app.utils.video_processor import video_processor
            from app.utils.video_transcoder import needs_transcoding, transcode_to_h264
            
            # 检查是否需要转码
            video_path = save_path
            if needs_transcoding(save_path):
                print(f"视频需要转码: {save_path}")
                try:
                    transcoded_path = transcode_to_h264(save_path)
                    # 更新 source_path 为转码后的文件
                    content.source_path = transcoded_path
                    video_path = transcoded_path
                    # 删除原始文件
                    os.remove(save_path)
                    print(f"视频转码完成: {transcoded_path}")
                except Exception as transcode_error:
                    print(f"视频转码失败，使用原始文件: {transcode_error}")
            
            result = await video_processor.process_video_with_task(
                video_path, str(content_id), ollama_client, task_manager, max_frames=8
            )
            content.description = result["overall_description"]
            content.content_metadata["duration"] = result["duration"]
            content.content_metadata["frame_count"] = len(result["frames"])
            
            # 生成向量
            if content.description:
                embeddings = await ollama_client.embed([content.description])
                content.embedding = embeddings[0] if embeddings else None
        except Exception as e:
            content.description = f"分析失败: {str(e)}"
    
    db.add(content)
    db.commit()
    
    return APIResponse.success(
        data={"id": str(content_id), "content_type": content_type.value},
        message=f"内容上传成功: {content_id}"
    )


@router.get("/")
def list_contents(
    content_type: Optional[str] = None,
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db)
):
    """列出所有内容"""
    query = db.query(Content)
    if content_type:
        query = query.filter(Content.content_type == content_type)
    
    contents = query.order_by(Content.created_at.desc()).offset(skip).limit(limit).all()
    return APIResponse.success(
        data=[c.to_dict() for c in contents],
        message="获取内容列表成功"
    )


@router.post("/search")
async def search_contents(
    query: str,
    content_type: Optional[str] = None,
    top_k: int = 5,
    db: Session = Depends(get_db)
):
    """语义搜索内容"""
    from sqlalchemy import text
    
    # 生成查询向量
    embeddings = await ollama_client.embed([query])
    query_embedding = embeddings[0]
    
    # 构建SQL
    sql = """
    SELECT 
        id,
        content_type,
        original_name,
        extracted_text,
        description,
        1 - (embedding <=> :query_embedding) as similarity
    FROM contents
    WHERE 1 - (embedding <=> :query_embedding) > 0.5
    """
    
    params = {"query_embedding": str(query_embedding), "top_k": top_k}
    
    if content_type:
        # 转换为大写以匹配PostgreSQL枚举值
        sql += " AND content_type = :content_type"
        params["content_type"] = content_type.upper()
    
    sql += " ORDER BY embedding <=> :query_embedding LIMIT :top_k"
    
    result = db.execute(text(sql), params)
    
    results = []
    for row in result:
        results.append({
            "id": str(row.id),
            "content_type": row.content_type,
            "original_name": row.original_name,
            "text": row.extracted_text or row.description,
            "similarity": row.similarity
        })
    
    return APIResponse.success(
        data={"results": results, "total": len(results)},
        message="搜索完成"
    )


@router.get("/{content_id}")
def get_content_detail(content_id: str, db: Session = Depends(get_db)):
    """获取内容详情"""
    content = db.query(Content).filter(Content.id == content_id).first()
    if not content:
        raise HTTPException(status_code=404, detail="内容不存在")
    
    # 构建URL
    url = None
    if content.content_type.value == "VIDEO":
        url = f"/uploads/contents/{os.path.basename(content.source_path)}"
    
    return APIResponse.success(
        data={
            "id": str(content.id),
            "url": url,
            "type": content.content_type.value.lower(),
            "original_name": content.original_name,
            "description": content.description,
            "extracted_text": content.extracted_text,
            "metadata": content.content_metadata,
            "created_at": content.created_at.isoformat() if content.created_at else None
        },
        message="获取内容详情成功"
    )


@router.delete("/{content_id}")
def delete_content(content_id: str, db: Session = Depends(get_db)):
    """删除内容"""
    content = db.query(Content).filter(Content.id == content_id).first()
    if not content:
        raise HTTPException(status_code=404, detail="内容不存在")
    
    # 删除文件
    if os.path.exists(content.source_path):
        os.remove(content.source_path)
    
    db.delete(content)
    db.commit()
    
    return APIResponse.success(message="内容已删除")


# ==================== 图片专用端点（向后兼容） ====================

@router.post("/images/upload")
async def upload_image(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """上传图片并生成描述（专用端点，向后兼容）"""
    # 检查文件类型
    if not file.content_type or not file.content_type.startswith('image/'):
        raise HTTPException(status_code=400, detail="只支持图片文件")
    
    # 生成UUID
    file_id = uuid.uuid4()
    file_ext = os.path.splitext(file.filename)[1]
    file_path = os.path.join(IMAGE_UPLOAD_DIR, f"{file_id}{file_ext}")
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # 生成图片描述（使用视觉模型）
    try:
        description = await ollama_client.describe_image(file_path)
    except Exception:
        description = "无法生成图片描述"
    
    # 生成描述向量
    try:
        embeddings = await ollama_client.embed([description])
        embedding = embeddings[0] if embeddings else None
    except:
        embedding = None
    
    # 保存到数据库
    image = Image(
        id=file_id,
        image_path=file_path,
        description=description,
        description_embedding=embedding
    )
    db.add(image)
    db.commit()
    
    return APIResponse.success(
        data={
            "id": str(file_id),
            "image_url": f"/uploads/images/{file_id}{file_ext}",
            "description": description
        },
        message="图片上传成功"
    )


@router.get("/images/list")
def list_images(
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db)
):
    """列出所有图片（专用端点，向后兼容）"""
    images = db.query(Image).offset(skip).limit(limit).all()
    
    return APIResponse.success(
        data=[
            {
                "id": str(img.id),
                "url": f"/uploads/images/{os.path.basename(img.image_path)}",
                "type": "image",
                "description": img.description,
                "created_at": img.created_at.isoformat() if img.created_at else None
            }
            for img in images
        ],
        message="获取图片列表成功"
    )


@router.get("/images/{image_id}")
def get_image_detail(
    image_id: str,
    db: Session = Depends(get_db)
):
    """获取图片详情（专用端点，向后兼容）"""
    img = db.query(Image).filter(Image.id == image_id).first()
    
    if not img:
        raise HTTPException(status_code=404, detail="图片不存在")
    
    return APIResponse.success(
        data={
            "id": str(img.id),
            "url": f"/uploads/images/{os.path.basename(img.image_path)}",
            "type": "image",
            "description": img.description,
            "metadata": img.image_metadata,
            "created_at": img.created_at.isoformat() if img.created_at else None
        },
        message="获取图片详情成功"
    )


@router.post("/images/search")
async def search_images(
    query: str,
    top_k: int = 5,
    db: Session = Depends(get_db)
):
    """语义搜索图片（专用端点，向后兼容）"""
    # 生成查询向量
    embeddings = await ollama_client.embed([query])
    query_embedding = embeddings[0]
    
    # 执行向量搜索
    from sqlalchemy import text
    sql = """
    SELECT 
        id,
        image_path,
        description,
        1 - (description_embedding <=> :query_embedding) as similarity
    FROM images
    WHERE 1 - (description_embedding <=> :query_embedding) > 0.5
    ORDER BY description_embedding <=> :query_embedding
    LIMIT :top_k
    """
    
    result = db.execute(
        text(sql),
        {
            "query_embedding": str(query_embedding),
            "top_k": top_k
        }
    )
    
    images = []
    for row in result:
        images.append({
            "id": str(row.id),
            "image_url": f"/uploads/images/{os.path.basename(row.image_path)}",
            "description": row.description,
            "similarity": row.similarity
        })
    
    return APIResponse.success(
        data={"results": images, "total": len(images)},
        message="图片搜索完成"
    )


@router.delete("/images/{image_id}")
def delete_image(
    image_id: str,
    db: Session = Depends(get_db)
):
    """删除图片（专用端点，向后兼容）"""
    image = db.query(Image).filter(Image.id == image_id).first()
    
    if not image:
        raise HTTPException(status_code=404, detail="图片不存在")
    
    # 删除文件
    if os.path.exists(image.image_path):
        try:
            os.remove(image.image_path)
        except Exception as e:
            print(f"删除图片文件失败: {e}")
    
    db.delete(image)
    db.commit()
    
    return APIResponse.success(message="图片删除成功")


# ==================== 文档专用端点（向后兼容） ====================

@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """上传文件并创建文档（专用端点，向后兼容）"""
    # 解析文件内容
    content = await parse_file(file)
    
    # 确定文档类型
    doc_type = "text"
    if file.filename.endswith('.pdf'):
        doc_type = "pdf"
    elif file.filename.endswith(('.docx', '.doc')):
        doc_type = "docx"
    
    doc_data = DocumentCreate(
        title=file.filename,
        content=content,
        doc_type=doc_type,
        metadata={"original_filename": file.filename}
    )
    
    service = DocumentService(db)
    result = await service.create_document(doc_data)
    
    return APIResponse.success(
        data=result,
        message="文档上传成功"
    )


@router.get("/documents/list")
def list_documents(
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db)
):
    """列出所有文档（专用端点，向后兼容）"""
    service = DocumentService(db)
    results = service.list_documents(skip=skip, limit=limit)
    
    # 将SQLAlchemy模型转换为字典
    documents = []
    for doc in results:
        documents.append({
            "id": str(doc.id),
            "title": doc.title,
            "content": doc.content[:500] + "..." if len(doc.content) > 500 else doc.content,
            "doc_type": doc.doc_type,
            "metadata": doc.doc_metadata,
            "created_at": doc.created_at.isoformat() if doc.created_at else None
        })
    
    return APIResponse.success(
        data=documents,
        message="获取文档列表成功"
    )


@router.get("/documents/{doc_id}")
def get_document(
    doc_id: UUID,
    db: Session = Depends(get_db)
):
    """获取单个文档（专用端点，向后兼容）"""
    service = DocumentService(db)
    doc = service.get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    # 将SQLAlchemy模型转换为字典
    document = {
        "id": str(doc.id),
        "title": doc.title,
        "content": doc.content,
        "doc_type": doc.doc_type,
        "metadata": doc.doc_metadata,
        "created_at": doc.created_at.isoformat() if doc.created_at else None
    }
    
    return APIResponse.success(
        data=document,
        message="获取文档详情成功"
    )


@router.delete("/documents/{doc_id}")
def delete_document(
    doc_id: UUID,
    db: Session = Depends(get_db)
):
    """删除文档（专用端点，向后兼容）"""
    service = DocumentService(db)
    success = service.delete_document(doc_id)
    if not success:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    return APIResponse.success(message="文档删除成功")


@router.post("/documents/search")
async def search_documents(
    request: SearchRequest,
    db: Session = Depends(get_db)
):
    """语义搜索文档（专用端点，向后兼容）"""
    service = DocumentService(db)
    results = await service.search_similar(
        query=request.query,
        top_k=request.top_k
    )
    
    return APIResponse.success(
        data={
            "text_results": results,
            "total_results": len(results)
        },
        message="文档搜索完成"
    )


@router.post("/documents/create")
async def create_document(
    doc_data: DocumentCreate,
    db: Session = Depends(get_db)
):
    """创建文档（专用端点，向后兼容）"""
    service = DocumentService(db)
    doc = await service.create_document(doc_data)
    
    # 将SQLAlchemy模型转换为字典
    document = {
        "id": str(doc.id),
        "title": doc.title,
        "content": doc.content,
        "doc_type": doc.doc_type,
        "metadata": doc.doc_metadata,
        "created_at": doc.created_at.isoformat() if doc.created_at else None
    }
    
    return APIResponse.success(
        data=document,
        message="文档创建成功"
    )
